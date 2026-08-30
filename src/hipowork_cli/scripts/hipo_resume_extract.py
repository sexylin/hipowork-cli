"""简历文本提取：PDF/DOCX/TXT → 纯文本，供 LLM 解析或人工确认。

只负责文本抽取（轻依赖 pymupdf），不做 LLM 结构化 — 结构化由
hipo_resume_import.py 交给 AI 服务 / 用户自己的 LLM 完成。

用法：
  python3 hipo_resume_extract.py resume.pdf --out resume.txt
  python3 hipo_resume_extract.py resume.docx          # 打印到 stdout
  python3 hipo_resume_extract.py resume.txt
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from hipo_auth import require_py310


def extract_text(path: Path) -> str:
    """按扩展名提取纯文本。"""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in (".docx", ".doc"):
        return _extract_docx(path)
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"不支持的简历格式: {suffix}（支持 pdf/docx/txt/md）")


def _extract_pdf(path: Path) -> str:
    try:
        import pymupdf  # 新版 API；旧版可 import fitz
    except ImportError:  # pragma: no cover
        try:
            import fitz as pymupdf  # type: ignore[no-redef]
        except ImportError as exc:
            raise ImportError(
                "需要 pymupdf：pip install -r requirements-resume.txt"
            ) from exc
    doc = pymupdf.open(path)
    try:
        return "\n".join(page.get_text() for page in doc).strip()
    finally:
        doc.close()


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "需要 python-docx：pip install python-docx"
        ) from exc
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts).strip()


def main() -> int:
    require_py310()
    parser = argparse.ArgumentParser(description="简历文本提取（PDF/DOCX/TXT → 纯文本）")
    parser.add_argument("resume", help="简历文件路径")
    parser.add_argument("--out", default="", help="输出文件路径（默认打印到 stdout）")
    args = parser.parse_args()

    path = Path(args.resume)
    if not path.exists():
        print(f"❌ 文件不存在: {path}", file=sys.stderr)
        return 1

    try:
        text = extract_text(path)
    except (ValueError, ImportError) as exc:
        print(f"❌ {exc}", file=sys.stderr)
        return 1

    if not text or len(text.strip()) < 20:
        print("⚠️  提取文本为空或过短，简历可能是扫描件（图片型 PDF 需要 OCR）。", file=sys.stderr)

    if args.out:
        out = Path(args.out)
        out.write_text(text, encoding="utf-8")
        print(f"✅ 已提取 {len(text)} 字符 → {out}")
        return 0
    print(text)
    return 0


if __name__ == "__main__":
    sys.exit(main())